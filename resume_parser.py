import os
import re
import json
import ast
import requests
import pdfplumber
import docx
import hashlib
from settings import DATA_DIR, load_settings

RESUME_DATA_FILE = os.path.join(DATA_DIR, "resume_data.json")
RESUME_CACHE_FILE = os.path.join(DATA_DIR, "resume_cache.json")
CACHE_VERSION = "v2"

# In-memory cache registry to serve results in sub-milliseconds
_IN_MEMORY_CACHE = {}

# Reuse TCP connections via a global session object for speed and efficiency
session = requests.Session()

def calculate_file_hash(file_path):
    """Calculate MD5 hash of a file."""
    hasher = hashlib.md5()
    try:
        with open(file_path, 'rb') as f:
            buf = f.read(65536)
            while len(buf) > 0:
                hasher.update(buf)
                buf = f.read(65536)
        return hasher.hexdigest()
    except Exception as e:
        print(f"Error calculating file hash: {e}")
        return None


def extract_text_from_file(file_path):
    """Extract text from .pdf or .docx files."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        text_list = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_list.append(text)
        return '\n'.join(text_list)
    elif ext == '.docx':
        doc = docx.Document(file_path)
        full_text = []
        
        # 1. Extract from headers of all sections
        try:
            for section in doc.sections:
                header = getattr(section, 'header', None)
                if header:
                    for para in header.paragraphs:
                        val = para.text.strip()
                        if val:
                            full_text.append(val)
        except Exception as e:
            print(f"Warning: could not extract headers from docx: {e}")
            
        # 2. Extract from main paragraphs
        for para in doc.paragraphs:
            val = para.text.strip()
            if val:
                full_text.append(val)
                
        # 3. Extract from tables (including paragraphs inside table cells)
        for table in doc.tables:
            for row in table.rows:
                row_cells_text = []
                for cell in row.cells:
                    cell_text = []
                    for para in cell.paragraphs:
                        val = para.text.strip()
                        if val:
                            cell_text.append(val)
                    cell_val = "\n".join(cell_text) if cell_text else cell.text.strip()
                    if cell_val:
                        row_cells_text.append(cell_val)
                if row_cells_text:
                    full_text.append(" | ".join(row_cells_text))
                    
        # 4. Extract from footers of all sections
        try:
            for section in doc.sections:
                footer = getattr(section, 'footer', None)
                if footer:
                    for para in footer.paragraphs:
                        val = para.text.strip()
                        if val:
                            full_text.append(val)
        except Exception as e:
            print(f"Warning: could not extract footers from docx: {e}")
            
        return '\n'.join(full_text)
    else:
        raise ValueError("Unsupported file format. Please upload .pdf or .docx")


def extract_json_from_response(text):
    """Extract and parse JSON from API response, removing markdown code blocks if present."""
    match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', text)
    if match:
        text = match.group(1)
    
    text = text.strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        # Fallback: locate first '{' and last '}'
        start = text.find('{')
        end = text.rfind('}')
        if start != -1 and end != -1:
            snippet = text[start:end+1]
            try:
                return json.loads(snippet)
            except json.JSONDecodeError:
                # Fallback to ast.literal_eval for single-quoted Python dict representation
                try:
                    parsed = ast.literal_eval(snippet)
                    if isinstance(parsed, dict):
                        return parsed
                except Exception:
                    pass
        raise ValueError(f"Could not parse valid JSON from AI response: {text}")

def parse_resume_with_gemini(text, api_key):
    """Send resume text to Gemini API for parsing."""
    schema = (
        "{\n"
        "  \"name\": \"Full Name (string)\",\n"
        "  \"email\": \"Email address (string or null)\",\n"
        "  \"phone\": \"Phone number (string or null)\",\n"
        "  \"location\": \"Current location (string or null)\",\n"
        "  \"links\": {\n"
        "    \"github\": \"GitHub link/profile (string or null)\",\n"
        "    \"linkedin\": \"LinkedIn link/profile (string or null)\",\n"
        "    \"portfolio\": \"Portfolio link (string or null)\",\n"
        "    \"other\": [\"List of any other URLs/links found in the resume (strings)\"]\n"
        "  },\n"
        "  \"professional_summary\": \"Professional summary or objective (string or null)\",\n"
        "  \"total_years_experience\": Total number of years of experience (float, e.g. 3.4),\n"
        "  \"skills\": [\"Flat list of all technical and soft skills (strings)\"],\n"
        "  \"job_titles\": [\"Flat list of all historical job titles (strings)\"],\n"
        "  \"target_roles\": [\"Comprehensive list of ALL possible target job titles, roles, designations, and adjacent roles (at least 15-20 distinct variations, including junior/senior/specialist levels, tool-specific roles, and adjacent domain roles) the candidate is qualified to apply to, ensuring no job opportunities are missed. (strings)\"],\n"
        "  \"industries\": [\"Flat list of industries relevant to the experience (strings)\"],\n"
        "  \"work_experience\": [\n"
        "    {\n"
        "      \"job_title\": \"Job title (string)\",\n"
        "      \"company\": \"Company name (string)\",\n"
        "      \"location\": \"Job location (string or null)\",\n"
        "      \"start_date\": \"Start date (string)\",\n"
        "      \"end_date\": \"End date (string, or 'Present')\",\n"
        "      \"responsibilities\": [\"List of responsibilities, achievements, and bullet points (strings)\"]\n"
        "    }\n"
        "  ],\n"
        "  \"projects\": [\n"
        "    {\n"
        "      \"title\": \"Project title (string)\",\n"
        "      \"technologies\": [\"List of technologies used in the project (strings)\"],\n"
        "      \"description\": [\"Bullet points describing the project (strings)\"]\n"
        "    }\n"
        "  ],\n"
        "  \"education\": [\n"
        "    {\n"
        "      \"degree\": \"Degree title (string)\",\n"
        "      \"institution\": \"Institution/University name (string)\",\n"
        "      \"graduation_year\": \"Graduation year or range (string)\",\n"
        "      \"details\": \"GPA, percentage, CGPA, honors (string or null)\"\n"
        "    }\n"
        "  ],\n"
        "  \"certifications\": [\"List of certifications (strings)\"],\n"
        "  \"languages\": [\"List of spoken/written languages (strings)\"],\n"
        "  \"achievements\": [\"List of achievements/awards/extracurricular activities (strings)\"],\n"
        "  \"publications\": [\"List of publications/papers (strings)\"],\n"
        "  \"patents\": [\"List of patents (strings)\"],\n"
        "  \"volunteering\": [\n"
        "    {\n"
        "      \"role\": \"Volunteering role (string)\",\n"
        "      \"organization\": \"Organization name (string)\",\n"
        "      \"description\": [\"Bullet points describing volunteering (strings)\"]\n"
        "    }\n"
        "  ],\n"
        "  \"other_sections\": {\n"
        "    \"section_name\": \"Full content of any other section or miscellaneous info not matching standard keys (e.g. hobbies, interests, references, etc.)\"\n"
        "  }\n"
        "}"
    )

    headers = {"Content-Type": "application/json"}
    payload = {
        "contents": [{
            "parts": [{
                "text": f"Resume text to parse:\n{text}"
            }]
        }],
        "systemInstruction": {
            "parts": [{
                "text": (
                    "You are a professional high-speed resume parser. Parse the provided resume text and extract "
                    "all information into a structured JSON object according to the following schema. Ensure no details are lost. "
                    "For target_roles, brainstorm an exhaustive, comprehensive list of all possible job titles, designations, and adjacent roles (aim for at least 15-20 variations) that match the candidate's skills and experience. Do not limit to standard titles; include specialized, general, junior, senior, and related role titles so they do not miss any job postings. "
                    "Reply ONLY with the raw JSON object. Do not wrap the JSON in markdown formatting.\n\n"
                    f"Schema:\n{schema}"
                )
            }]
        },
        "generationConfig": {
            "responseMimeType": "application/json",
            "temperature": 0.1
        }
    }
    
    import time
    # Model fallback hierarchy for high reliability and speed
    models_to_try = ["gemini-2.5-pro", "gemini-2.5-flash", "gemini-2.5-flash-lite", "gemini-1.5-flash-latest", "gemini-flash-latest"]
    last_error = None
    
    for model_name in models_to_try:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={api_key}"
        
        for attempt in range(2):  # Try up to 2 attempts per model
            try:
                # Use global session object to reuse TCP connection (keep-alive)
                response = session.post(url, headers=headers, json=payload, timeout=25)
                if response.status_code == 429:
                    if attempt < 1:
                        time.sleep(3)
                        continue
                    else:
                        raise Exception(f"Gemini model {model_name} rate limit hit.")
                
                if response.status_code != 200:
                    try:
                        err_json = response.json()
                        err_msg = err_json.get("error", {}).get("message", response.text)
                    except Exception:
                        err_msg = response.text
                    
                    if response.status_code in (401, 403):
                        raise PermissionError(f"Gemini API Authentication/Permission Error ({response.status_code}): {err_msg}")
                    if response.status_code in (400, 404):
                        raise LookupError(f"Gemini API Bad Request/Not Found ({response.status_code}): {err_msg}")
                        
                    raise ValueError(f"Gemini API Error ({response.status_code}): {err_msg}")
                    
                resp_data = response.json()
                try:
                    ai_response_text = resp_data["candidates"][0]["content"]["parts"][0]["text"]
                except (KeyError, IndexError) as e:
                    raise ValueError("Invalid response format received from Gemini API.") from e
                    
                return extract_json_from_response(ai_response_text)
            except PermissionError as pe:
                print(f"Resume parser: Unrecoverable API permission error: {pe}")
                raise pe
            except LookupError as le:
                print(f"Resume parser: Model {model_name} is invalid/not found: {le}")
                last_error = le
                break  # Skip any subsequent attempts for this model, and move to the next model immediately
            except Exception as e:
                last_error = e
                # Fallback logging or debug print
                print(f"Resume parser: Attempt failed with model {model_name}: {e}")
                if attempt < 1:
                    time.sleep(1)
                
    raise last_error if last_error else ValueError("All Gemini parsing attempts failed.")


def parse_and_save_resume(file_path):
    """Main function to parse resume and save to resume_data.json."""
    settings = load_settings()
    api_key = settings.get("gemini_api_key")
    if not api_key:
        raise ValueError("Gemini API key is missing in settings. Please save your API key first.")
        
    current_hash = calculate_file_hash(file_path)
    
    # 1. In-Memory Cache Check: Return parsed results instantly if loaded during the current session
    if current_hash and current_hash in _IN_MEMORY_CACHE:
        cached_data = _IN_MEMORY_CACHE[current_hash]
        if cached_data.get("cache_version") == CACHE_VERSION:
            print("Resume parser: In-memory cache match found! Returning parsed results instantly.")
            try:
                with open(RESUME_DATA_FILE, "w", encoding="utf-8") as f:
                    json.dump(cached_data, f, indent=2)
            except Exception as write_err:
                print(f"Resume parser: Error syncing in-memory cache to active resume file: {write_err}")
            return cached_data
            
    # 2. Active Cache Check: Return cached parse results instantly if current file is the active one on disk
    if current_hash and os.path.exists(RESUME_DATA_FILE):
        try:
            if os.path.getsize(RESUME_DATA_FILE) > 0:
                with open(RESUME_DATA_FILE, "r", encoding="utf-8") as f:
                    cached_data = json.load(f)
                if cached_data.get("file_hash") == current_hash and cached_data.get("cache_version") == CACHE_VERSION:
                    print("Resume parser: Active resume cache match found! Returning parsed results instantly.")
                    _IN_MEMORY_CACHE[current_hash] = cached_data
                    return cached_data
        except Exception as cache_err:
            print(f"Resume parser: Active cache read error (ignoring cache): {cache_err}")
            
    # 3. Persistent Multi-Resume Cache Check: Look up hash registry of all previously parsed resumes
    if current_hash and os.path.exists(RESUME_CACHE_FILE):
        try:
            if os.path.getsize(RESUME_CACHE_FILE) > 0:
                with open(RESUME_CACHE_FILE, "r", encoding="utf-8") as f:
                    full_cache = json.load(f)
                if isinstance(full_cache, dict) and current_hash in full_cache:
                    cached_data = full_cache[current_hash]
                    if cached_data.get("cache_version") == CACHE_VERSION:
                        print("Resume parser: Persistent cache match found! Returning parsed results instantly.")
                        _IN_MEMORY_CACHE[current_hash] = cached_data
                        # Sync to the active resume data file
                        with open(RESUME_DATA_FILE, "w", encoding="utf-8") as f:
                            json.dump(cached_data, f, indent=2)
                        return cached_data
        except Exception as cache_err:
            print(f"Resume parser: Persistent cache read error (ignoring cache): {cache_err}")
            
    text = extract_text_from_file(file_path)
    if not text.strip():
        raise ValueError("Could not extract any text from the uploaded resume file.")
        
    parsed_data = parse_resume_with_gemini(text, api_key)
    
    # Ensure raw text is stored so we do not miss any information
    parsed_data["full_extracted_text"] = text
    parsed_data["cache_version"] = CACHE_VERSION
    
    # Store the file hash to enable instant caching on future uploads
    if current_hash:
        parsed_data["file_hash"] = current_hash
        
    # Update in-memory registry
    if current_hash:
        _IN_MEMORY_CACHE[current_hash] = parsed_data
        
    # Update persistent multi-resume cache file
    if current_hash:
        try:
            full_cache = {}
            if os.path.exists(RESUME_CACHE_FILE) and os.path.getsize(RESUME_CACHE_FILE) > 0:
                with open(RESUME_CACHE_FILE, "r", encoding="utf-8") as f:
                    full_cache = json.load(f)
                if not isinstance(full_cache, dict):
                    full_cache = {}
            full_cache[current_hash] = parsed_data
            with open(RESUME_CACHE_FILE, "w", encoding="utf-8") as f:
                json.dump(full_cache, f, indent=2)
        except Exception as write_err:
            print(f"Resume parser: Error writing to persistent cache file: {write_err}")
    
    # Save parsed data to active resume file
    with open(RESUME_DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(parsed_data, f, indent=2)
        
    return parsed_data


def clear_cache():
    """Clear both persistent and in-memory caches."""
    _IN_MEMORY_CACHE.clear()
    if os.path.exists(RESUME_CACHE_FILE):
        try:
            os.remove(RESUME_CACHE_FILE)
        except Exception as e:
            print(f"Error clearing cache file: {e}")

