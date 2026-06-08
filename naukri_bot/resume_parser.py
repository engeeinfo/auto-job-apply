import os
import json
from typing import Dict, Any, Optional
from pypdf import PdfReader
import docx

class ResumeParser:
    """
    Extracts text from PDF/Word/TXT resumes and structures them into standardized JSON
    via the AI Engine, storing a local cache to avoid duplicate API calls.
    """
    def __init__(self, data_path: str = None):
        if data_path is None:
            # Default to naukri_bot/data/resume_data.json relative to this file
            base_dir = os.path.dirname(os.path.abspath(__file__))
            self.data_path = os.path.join(base_dir, "data", "resume_data.json")
        else:
            self.data_path = os.path.abspath(data_path)

    def load_cached_resume(self) -> Optional[Dict[str, Any]]:
        """
        Retrieves cached structured resume data if it exists.
        """
        if os.path.exists(self.data_path):
            try:
                with open(self.data_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if isinstance(data, dict):
                        return data
            except Exception as e:
                print(f"[ResumeParser] Error reading cached resume: {e}")
        return None

    def save_resume_data(self, data: Dict[str, Any]) -> None:
        """
        Saves structured resume data to local JSON cache.
        """
        os.makedirs(os.path.dirname(self.data_path), exist_ok=True)
        try:
            with open(self.data_path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2)
        except Exception as e:
            print(f"[ResumeParser] Error caching resume: {e}")

    def extract_text(self, file_path: str) -> str:
        """
        Reads raw text content from PDF, DOCX, or TXT formats.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Resume file not found at: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        text_parts = []

        if ext == ".pdf":
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            raw_text = "\n".join(text_parts)

        elif ext == ".docx":
            doc = docx.Document(file_path)
            # Pull paragraph text
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # Pull table cell text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            raw_text = "\n".join(text_parts)

        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                raw_text = f.read()
        else:
            raise ValueError(f"Unsupported resume file type: {ext}. Use PDF, DOCX, or TXT.")

        cleaned_text = raw_text.strip()
        if not cleaned_text:
            raise ValueError("Extracted resume text is empty.")
            
        return cleaned_text

    async def parse_resume(self, file_path: str, ai_engine) -> Dict[str, Any]:
        """
        Extracts text, structures it using the AI engine, caches the JSON, and returns it.
        """
        raw_text = self.extract_text(file_path)
        structured_data = await ai_engine.parse_resume_text(raw_text)
        
        # Save filepath metadata inside the structure for verification/UI display
        structured_data["_file_path"] = os.path.abspath(file_path)
        self.save_resume_data(structured_data)
        return structured_data
